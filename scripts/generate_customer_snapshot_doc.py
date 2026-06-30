#!/usr/bin/env python3
"""Generate customer-facing POC snapshot as .docx from .txt content."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
TXT = ROOT / "docs" / "customer" / "HDFC_Eligibility_Engine_POC_Snapshot.txt"
DOCX = ROOT / "docs" / "customer" / "HDFC_Eligibility_Engine_POC_Snapshot.docx"


def main() -> None:
    text = TXT.read_text(encoding="utf-8")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_paragraph("HDFC Eligibility Engine — Mock Producer API (POC)")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Customer Snapshot")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True

    doc.add_paragraph()

    for line in text.splitlines():
        stripped = line.strip()
        if not stripped or stripped.startswith("="):
            continue
        if stripped[0].isdigit() and "." in stripped[:4]:
            p = doc.add_paragraph(stripped)
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(12)
        elif stripped.endswith("---") or stripped.endswith("..."):
            doc.add_paragraph(stripped.rstrip("-"))
        else:
            doc.add_paragraph(line)

    DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX)
    print(f"Wrote {DOCX}")


if __name__ == "__main__":
    main()
